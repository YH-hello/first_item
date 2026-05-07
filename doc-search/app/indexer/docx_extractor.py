"""
Word(DOCX) 텍스트 및 페이지 정보 추출기
라이브러리: python-docx
페이지 감지: XML의 lastRenderedPageBreak / w:pageBreak 태그 활용
딥링크: win32com Word.Application → Selection.GoTo(wdGoToPage)
"""
import logging
from typing import Dict, List

import docx
from lxml import etree

from .base_extractor import BaseExtractor, PageChunk

logger = logging.getLogger(__name__)

# Word XML 네임스페이스
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _has_page_break(paragraph) -> bool:
    """단락 XML에서 명시적 페이지 구분 여부 감지"""
    xml_str = etree.tostring(paragraph._p, encoding="unicode")
    return "lastRenderedPageBreak" in xml_str or "w:pageBreak" in xml_str


class DOCXExtractor(BaseExtractor):
    """python-docx를 이용한 페이지 단위 텍스트 추출"""

    def extract(self) -> List[PageChunk]:
        chunks: List[PageChunk] = []
        try:
            doc = docx.Document(str(self.file_path))
            page_texts: Dict[int, List[str]] = {}
            current_page = 1

            for para in doc.paragraphs:
                if _has_page_break(para):
                    current_page += 1

                text = para.text.strip()
                if text:
                    page_texts.setdefault(current_page, []).append(text)

            # 표(table) 셀 텍스트도 수집
            for table in doc.tables:
                for row in table.rows:
                    row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_parts:
                        page_texts.setdefault(current_page, []).append(" | ".join(row_parts))

            abs_path = str(self.file_path.resolve()).replace("\\", "\\\\")

            for page_num, texts in sorted(page_texts.items()):
                text = "\n".join(texts)
                if not text.strip():
                    continue

                page_label = f"{page_num}페이지"
                # wdGoToPage=1, wdGoToAbsolute=1
                page_link_cmd = (
                    f"python -c \""
                    f"import win32com.client; "
                    f"word=win32com.client.Dispatch('Word.Application'); "
                    f"word.Visible=True; "
                    f"doc=word.Documents.Open(r'{abs_path}'); "
                    f"doc.ActiveWindow.Selection.GoTo(What=1, Which=1, Count={page_num})\""
                )

                chunks.append(
                    self._build_chunk(
                        text=text,
                        page_num=page_num,
                        page_label=page_label,
                        page_link_cmd=page_link_cmd,
                    )
                )

        except Exception as exc:
            logger.error("[DOCX] 추출 오류 %s: %s", self.file_path, exc, exc_info=True)

        return chunks
